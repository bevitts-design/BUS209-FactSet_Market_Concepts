from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
import html

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)

NAVY = "12304A"; BLUE = "176B87"; TEAL = "1B998B"; GOLD = "E6A817"
PALE = "EAF5F4"; SKY = "EAF2F8"; INK = "243746"; MUTED = "5B6B78"; WHITE = "FFFFFF"

overview = [
    ("Course number", "BUS209"), ("Course title", "FactSet Financial Concepts (FFC)"),
    ("Credits", "1"), ("Class type", "Lecture / discussion / presentation / FactSet certification videos"),
    ("Semester", "Fall 2026"), ("Faculty", "Bethany Evitts, CFA — Assistant Professor"),
    ("Office", "Gerrish 356"), ("Telephone", "617-877-2001 (text accessible)"),
    ("Email", "bevitts@endicott.edu"), ("Office hours", "See electronic calendar"),
]

catalog = ("FactSet Financial Concepts (FFC) introduces the characteristics of four core areas of finance. "
           "The core areas covered are Economic Indicators, Currencies, Fixed Income, and Equities. The sections "
           "are woven together from FactSet data, news, and analytics. Students, individually or in small groups, "
           "will make class presentations using the FactSet terminal on at least one topic covered in each section.")

outcomes = [
    "Explain how investors use economic indicators to gauge the health of the economy.",
    "Explain the qualities of good economic indicators.",
    "Identify the three main drivers of currency valuation.",
    "Demonstrate how investors and businesses are affected by currency markets and how they manage currency risk.",
    "Describe how yields facilitate comparison across the vast diversity of the bond market.",
    "Interpret the meaning of the four major shifts in the yield curve and the meaning of curve inversion.",
    "Calculate equity index performance from the performance of single stocks.",
    "Identify the three types of relative valuation and the role of future earnings growth when assessing fair value.",
]

assessments = [
    ("Country Economic Indicators & Markets", "20%"),
    ("Equities / Pitch Book &/or Credit Presentation", "20%"),
    ("Final Presentation", "20%"),
    ("FactSet Certifications", "20%"),
    ("Attendance / Participation", "20%"),
]

evaluation_summary = (
    "Student learning will be assessed through the following. Core section deliverables "
    "and presentations account for 60%; FactSet Certifications account for 20%; and "
    "logistics, timeliness, preparedness, attendance, and participation account for 20%."
)

grades = [("A","94–100"),("A−","90–93"),("B+","87–89"),("B","84–86"),("B−","80–83"),
          ("C+","77–79"),("C","74–76"),("C−","70–73"),("D+","67–69"),("D","64–66"),
          ("D−","60–63"),("F","< 60"),("WX","Withdrawn Failed")]

schedule = [
    (1,"September 14","Introductions & FactSet Certification Expectations; FactSet login; Company Scavenger Hunt","class"),
    (2,"September 21","Exploring the Markets with FactSet; Economic Indicators; assignment of Economic Indicators topics","class"),
    (3,"September 28","Economic Indicator Discussion; Wealth Management: News & Market Performance","class"),
    (4,"October 5","Economic Indicator Presentation","presentation"),
    (5,"October 12","No Class — Indigenous Peoples/Columbus Day (observed)","break"),
    (6,"October 19","Equity Market; Equity Research; Accounting and Company Information","class"),
    (7,"October 26","Equity Market Industry Analysis and ESG","class"),
    (8,"November 2","Equity Presentations","presentation"),
    (9,"November 9","Fixed Income — Introduction & Credit Analysis","class"),
    (10,"November 16","No Class (instructor-planned; confirm in Canvas)","break"),
    (11,"November 23","No Class — Thanksgiving recess","break"),
    (12,"November 30","Fixed Income Market Presentations","presentation"),
    (13,"December 7","FactSet Financial Concepts Final Paper and Presentation","presentation"),
    (14,"Finals week: December 14–18","FactSet Financial Concepts Certifications","final"),
]

policies = [
    ("Academic Integrity Policy", "Students are required to abide by the Academic Integrity Policy of Endicott College."),
    ("AI Policy", "Students are required to abide by the AI Policy of Endicott College, and violations are subject to the college’s Academic Integrity Policy. Instructor course policy to be finalized: include any course-wide policies, procedures for using and citing AI use, and the penalty for violations."),
    ("Attendance Policy", "Students are required to abide by the Attendance Policy of Endicott College. Attendance in this course is required, including the last day of scheduled classes in the semester and the final exam. Students who fail to attend 3 class meetings will be dismissed from this course. If you are absent due to an athletic contest, please have the appropriate paperwork completed in a timely fashion (before the game). You will not be permitted to miss class for workouts or practice."),
    ("Turnitin Policy", "By taking this course, students agree that all required assignments may be subject to submission for ‘similarity review’ to Turnitin.com, a tool intended not just to detect instances of plagiarism, but to prevent it as well. The tool is intended to help students identify passages that are unoriginal, incorrectly cited, or lacking appropriate source information. Submitted assignments may also be archived in the Turnitin.com database for checking possible future instances of plagiarism, additional similarity searches, and other educational purposes at the discretion of the instructor. For more information, please review the Privacy and Security guide at Turnitin.com."),
]

support = [
    ("Accessibility Services", "Endicott College provides equal educational opportunities for all students regardless of disability status. If you believe that you qualify as a person with a disability as defined by Section 504 of the Rehabilitation Act of 1973, the Americans with Disabilities Act (ADA) of 1990 and its Amendment Act of 2008 (ADAAA), you are encouraged to register with the Center for Accessibility Services Office to request accommodations, auxiliary aides, and/or support. The Center for Accessibility Services is located on the 2nd Floor of the Diane M. Halle Library. Please visit the Center for Accessibility Services website and contact access@endicott.edu with questions."),
    ("Academic Support", "The Division of Academic Success believes that every student can benefit from having a thought partner who supports their learning. We offer innovative and individualized services that motivate you as you pursue your educational goals."),
    ("Content and Writing Tutoring", "Work with content tutors who can help you understand, remember, and apply course content, or with writing tutors who will support your growth as writers and thinkers. Students can schedule free tutoring sessions on TutorTrac or stop by Halle Library 204."),
    ("Quick-Connect Coaching", "Meet with a professional academic coach for thirty minutes to devise a solution to an immediate challenge, including unpacking a difficult assignment, creating a specific study schedule for an upcoming assessment, prioritizing multiple responsibilities, and more. Students can schedule free quick-connect sessions on TutorTrac or stop by Halle Library 204."),
    ("Academic Coaching", "Grow your academic self-awareness and deepen your connection to your education by partnering once, twice, or three times a week with a professional academic coach to develop academic confidence, resilience, and grit; develop time-management and organizational systems and strategies; learn how to set and adjust goals; and recognize how best to use resources. This program has an additional cost. To enroll, email academiccoaching@endicott.edu."),
    ("Workshops", "Semester workshops help students build skills in crucial areas for learning and succeeding: time management, goal setting, note-taking strategies, and metacognition. See the website for a complete list of programs."),
    ("Title IX 2024 Policy", "Institutions cannot discriminate based on a student’s pregnancy or parental status—past or present—or any related conditions, medical or otherwise. Qualified students may be eligible for academic adjustments enabling them to continue their academic process. For more information about Title IX protections regarding pregnancy or parental status, contact Endicott’s Title IX Officer, Christy Galatis, at cgalatis@endicott.edu, (978) 998-7746, or www.endicott.edu/title-ix."),
]

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc.get_or_add_tcPr(); el = tc.first_child_found_in('w:tcMar')
    if el is None: el=OxmlElement('w:tcMar'); tc.append(el)
    for side,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        x=el.find(qn('w:'+side))
        if x is None: x=OxmlElement('w:'+side); el.append(x)
        x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa')

def set_font(run, size=10.5, bold=False, color=INK, name='Aptos'):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),name); run._element.rPr.rFonts.set(qn('w:hAnsi'),name)
    run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(color)

def set_repeat(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:tblHeader'); e.set(qn('w:val'),'true'); trPr.append(e)

def set_widths(table, widths):
    table.autofit=False
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(w)); grid.append(gc)
    pr=table._tbl.tblPr; tw=pr.find(qn('w:tblW')); tw.set(qn('w:type'),'dxa'); tw.set(qn('w:w'),str(sum(widths)))
    ind=pr.find(qn('w:tblInd'))
    if ind is None: ind=OxmlElement('w:tblInd'); pr.append(ind)
    ind.set(qn('w:type'),'dxa'); ind.set(qn('w:w'),'120')
    for row in table.rows:
        for c,w in zip(row.cells,widths):
            c.width=Inches(w/1440); c._tc.get_or_add_tcPr().tcW.set(qn('w:type'),'dxa'); c._tc.tcPr.tcW.set(qn('w:w'),str(w)); margins(c)

def add_bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); set_font(p.add_run(text)); return p

def add_heading(doc, text, level=1):
    p=doc.add_heading(text, level=level); p.paragraph_format.keep_with_next=True; return p

def build_docx():
    d=Document(); s=d.sections[0]; s.top_margin=Inches(.7); s.bottom_margin=Inches(.7); s.left_margin=Inches(.8); s.right_margin=Inches(.8); s.header_distance=Inches(.3); s.footer_distance=Inches(.35)
    styles=d.styles
    normal=styles['Normal']; normal.font.name='Aptos'; normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.08
    for name,size,color,before,after in [('Title',30,NAVY,0,6),('Subtitle',14,BLUE,0,12),('Heading 1',18,NAVY,16,7),('Heading 2',13,BLUE,10,4),('Heading 3',11,TEAL,8,3)]:
        st=styles[name]; st.font.name='Aptos Display' if name!='Normal' else 'Aptos'; st.font.size=Pt(size); st.font.bold=name!='Subtitle'; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    hdr=s.header.paragraphs[0]; hdr.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_font(hdr.add_run('BUS209  •  FALL 2026'),9,True,TEAL)
    foot=s.footer.paragraphs[0]; foot.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(foot.add_run('Endicott College  |  Curtis L. Gerrish School of Business'),8.5,False,MUTED)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(18); set_font(p.add_run('ENDICOTT COLLEGE'),11,True,TEAL)
    p=d.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('FactSet Financial Concepts')
    p=d.add_paragraph(style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('BUS209  •  Course Syllabus  •  Fall 2026')
    t=d.add_table(rows=5, cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER; set_widths(t,[1500,3180,1500,3180])
    for i,((l1,v1),(l2,v2)) in enumerate(zip(overview[:5],overview[5:])):
        for j,(lab,val) in enumerate([(l1,v1),(l2,v2)]):
            c=t.cell(i,j*2); shade(c, NAVY); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=c.paragraphs[0]; set_font(p.add_run(lab.upper()),8.5,True,WHITE)
            c=t.cell(i,j*2+1); shade(c, SKY if i%2==0 else WHITE); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=c.paragraphs[0]; set_font(p.add_run(val),9.5,False,INK)
    add_heading(d,'Course Overview'); d.add_paragraph(catalog)
    add_heading(d,'Learning Outcomes'); d.add_paragraph('Upon completion of this course, students will be able to:')
    for x in outcomes: add_bullet(d,x)
    add_heading(d,'Teaching & Learning Strategies'); d.add_paragraph('Course sessions include lecture, group discussions, and hands-on FactSet terminal exercises, both in class and for homework.')
    d.add_paragraph('In addition to the in-class sessions, it is your duty to regularly—at least once per day—check your email and Canvas to ensure you understand what you are expected to complete for this class.')
    d.add_paragraph('It is in your best interest to keep up with assignments. Please contact me with questions about the content. I am here to help. If you find any concepts unclear, let me know and let me help you. I want you to do your best in this course.')
    add_heading(d,'Readings & Materials'); add_bullet(d,'FactSet modules in video form.'); add_bullet(d,'FactSet Certificate programs, free for students who access learning information through the FactSet app/web pages.'); add_bullet(d,'Required: FactSet Desktop Application and Microsoft Office 365.')
    add_heading(d,'Evaluation Methods'); d.add_paragraph(evaluation_summary)
    t=d.add_table(rows=1,cols=2); set_widths(t,[7200,2160]); t.alignment=WD_TABLE_ALIGNMENT.CENTER; set_repeat(t.rows[0])
    for c,txt in zip(t.rows[0].cells,['Assessment','Weight']): shade(c,NAVY); set_font(c.paragraphs[0].add_run(txt),9.5,True,WHITE)
    for name,w in assessments:
        cells=t.add_row().cells
        for c,txt in zip(cells,[name,w]): set_font(c.paragraphs[0].add_run(txt),9.5,False,INK); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if len(t.rows)%2==0:
            for c in cells: shade(c,SKY)
    add_heading(d,'Grading'); d.add_paragraph('Each student will be assigned a letter grade as follows:')
    t=d.add_table(rows=1,cols=4); set_widths(t,[1400,3280,1400,3280]);
    pairs=[grades[i:i+2] for i in range(0,len(grades),2)]
    for c,txt in zip(t.rows[0].cells,['Grade','Range','Grade','Range']): shade(c,BLUE); set_font(c.paragraphs[0].add_run(txt),9,True,WHITE)
    for pair in pairs:
        row=t.add_row().cells
        vals=[]
        for g,rng in pair: vals += [g,rng]
        while len(vals)<4: vals += ['','']
        for c,txt in zip(row,vals): set_font(c.paragraphs[0].add_run(txt),9.5,False,INK)
    add_heading(d,'Course Policies');
    for title,text in policies: add_heading(d,title,2); d.add_paragraph(text)
    add_heading(d,'Student Support');
    for title,text in support: add_heading(d,title,2); d.add_paragraph(text)
    add_heading(d,'Course Expectations'); d.add_paragraph('For each credit hour, students are expected to spend a minimum of two hours on work outside of class each week. For this one-credit course, that means a minimum of two hours each week.')
    d.add_paragraph('Students must review the online Academic Calendar published by the Registrar’s Office: https://www.endicott.edu/academics/academic-resources-support/academic-calendar/undergraduate (the original syllabus also referenced http://tinyurl.com/hbmfywj).')
    d.add_paragraph('Class attendance is expected of all students up to and including the last day of scheduled finals in the semester. Students must plan accordingly.')
    add_heading(d,'Topic Outline & Timeline'); d.add_paragraph('The table below is a rough guide to the topics and activities in this course. Starting in Week 1, you have something due almost every week. Any schedule changes will be announced via Canvas; log in at least once per day to check announcements and keep pace with assignments.')
    t=d.add_table(rows=1,cols=3); set_widths(t,[750,1750,6860]); set_repeat(t.rows[0])
    for c,txt in zip(t.rows[0].cells,['Week','Week of','Assignments Due / Weekly Class Topics']): shade(c,NAVY); set_font(c.paragraphs[0].add_run(txt),9,True,WHITE)
    for wk,date,topic,kind in schedule:
        cells=t.add_row().cells; fill={'break':'FDF3D7','presentation':'EAF5F4','final':'EDE7F6'}.get(kind,WHITE)
        for c,txt in zip(cells,[str(wk),date,topic]): shade(c,fill); set_font(c.paragraphs[0].add_run(txt),9,False,INK); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_heading(d,'Subject to Change Statement'); d.add_paragraph('The syllabus sets forth my course objectives and my best estimate of what we will be able to cover during the semester. To tailor the course to the specific needs of the students, I might modify the syllabus during the semester; for this reason, it is possible that not all material will be covered and/or that additional material may be assigned.')
    path=OUT/'BUS209_FactSet_Financial_Concepts_Syllabus_Fall_2026.docx'; d.save(path); return path

def build_html():
    def esc(x): return html.escape(str(x))
    cards=''.join(f'<div class="fact"><span>{esc(k)}</span><strong>{esc(v)}</strong></div>' for k,v in overview)
    outcome_html=''.join(f'<li>{esc(x)}</li>' for x in outcomes)
    assess_html=''.join(f'<div class="weight"><span>{esc(n)}</span><strong>{esc(w)}</strong></div>' for n,w in assessments)
    grades_html=''.join(f'<div><b>{esc(g)}</b><span>{esc(r)}</span></div>' for g,r in grades)
    policy_html=''.join(f'<details open><summary>{esc(t)}</summary><p>{esc(x)}</p></details>' for t,x in policies)
    support_html=''.join(f'<details><summary>{esc(t)}</summary><p>{esc(x)}</p></details>' for t,x in support)
    schedule_html=''.join(f'<tr class="{kind}"><td>{wk}</td><td>{esc(date)}</td><td>{esc(topic)}</td></tr>' for wk,date,topic,kind in schedule)
    doc=f'''<!doctype html><html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>BUS209 | Fall 2026 Syllabus</title><style>
    :root{{--navy:#12304a;--blue:#176b87;--teal:#1b998b;--gold:#e6a817;--ink:#243746;--pale:#eaf5f4;--sky:#eaf2f8;--paper:#fff;--line:#d8e3e8}}*{{box-sizing:border-box}}html{{scroll-behavior:smooth}}body{{margin:0;background:#f4f7f8;color:var(--ink);font:16px/1.58 system-ui,-apple-system,"Segoe UI",sans-serif}}header{{background:linear-gradient(135deg,var(--navy),#18536a 65%,var(--teal));color:white;padding:4.5rem max(5vw,1.25rem) 3.8rem;position:relative;overflow:hidden}}header:after{{content:"";position:absolute;width:24rem;height:24rem;border-radius:50%;right:-7rem;top:-11rem;background:#ffffff14}}.kicker{{text-transform:uppercase;letter-spacing:.16em;font-weight:800;color:#a8eee5}}h1{{font-size:clamp(2.3rem,6vw,4.7rem);line-height:1.02;margin:.5rem 0}}header p{{font-size:1.2rem;max-width:50rem}}nav{{position:sticky;top:0;z-index:5;background:#fff;box-shadow:0 2px 12px #12304a18;padding:.7rem max(5vw,1rem);display:flex;gap:.5rem;overflow:auto}}nav a{{white-space:nowrap;color:var(--navy);text-decoration:none;font-weight:700;padding:.45rem .75rem;border-radius:99px}}nav a:hover,nav a:focus{{background:var(--pale)}}main{{width:min(1120px,92vw);margin:2rem auto 5rem}}section{{background:white;border:1px solid var(--line);border-radius:18px;padding:clamp(1.25rem,3vw,2.3rem);margin:1.2rem 0;box-shadow:0 8px 25px #12304a0d}}h2{{color:var(--navy);font-size:2rem;margin:.1rem 0 1rem}}h3{{color:var(--blue);margin-top:1.5rem}}.facts{{display:grid;grid-template-columns:repeat(auto-fit,minmax(230px,1fr));gap:.8rem}}.fact{{background:var(--sky);border-radius:12px;padding:1rem}}.fact span{{display:block;font-size:.76rem;text-transform:uppercase;letter-spacing:.08em;color:var(--blue);font-weight:800}}.fact strong{{display:block;margin-top:.2rem}}.callout{{border-left:6px solid var(--gold);background:#fff8e6;padding:1rem 1.2rem;border-radius:0 12px 12px 0}}li{{margin:.45rem 0}}.weights{{display:grid;grid-template-columns:repeat(auto-fit,minmax(220px,1fr));gap:.8rem}}.weight{{padding:1rem;border-radius:12px;background:var(--pale);display:flex;justify-content:space-between;gap:1rem}}.weight strong{{font-size:1.4rem;color:var(--teal)}}.grades{{display:grid;grid-template-columns:repeat(auto-fit,minmax(110px,1fr));gap:.55rem}}.grades div{{display:flex;justify-content:space-between;background:var(--sky);padding:.65rem .8rem;border-radius:8px}}details{{border-top:1px solid var(--line);padding:.75rem 0}}summary{{font-weight:800;color:var(--blue);cursor:pointer}}table{{width:100%;border-collapse:separate;border-spacing:0;font-size:.95rem}}th{{background:var(--navy);color:white;text-align:left;padding:.8rem}}th:first-child{{border-radius:10px 0 0 0}}th:last-child{{border-radius:0 10px 0 0}}td{{padding:.8rem;border-bottom:1px solid var(--line);vertical-align:top}}tr.break td{{background:#fff8e6}}tr.presentation td{{background:var(--pale)}}tr.final td{{background:#f2ecff}}footer{{text-align:center;color:#62727d;padding:2rem}}a{{color:var(--blue)}}@media(max-width:700px){{table{{display:block;overflow-x:auto}}nav{{font-size:.86rem}}section{{border-radius:12px}}}}@media print{{nav{{display:none}}body{{background:white;font-size:10pt}}header{{padding:1.5rem;background:var(--navy)!important;-webkit-print-color-adjust:exact;print-color-adjust:exact}}header h1{{font-size:28pt}}main{{width:100%;margin:0}}section{{box-shadow:none;break-inside:auto;margin:.5rem 0;padding:1rem}}details{{break-inside:avoid}}details>*{{display:block}}tr{{break-inside:avoid}}}}
    </style></head><body><header><div class="kicker">Endicott College · Curtis L. Gerrish School of Business</div><h1>FactSet Financial Concepts</h1><p>BUS209 · Fall 2026 · A one-credit path through economic indicators, currencies, fixed income, equities, and FactSet-powered presentation skills.</p></header><nav aria-label="Syllabus sections"><a href="#overview">Overview</a><a href="#outcomes">Outcomes</a><a href="#assessment">Assessment</a><a href="#policies">Policies</a><a href="#support">Support</a><a href="#schedule">Schedule</a></nav><main>
    <section id="overview"><h2>Course Overview</h2><div class="facts">{cards}</div><h3>Catalog Description</h3><p>{esc(catalog)}</p></section>
    <section id="outcomes"><h2>What You’ll Learn</h2><ol>{outcome_html}</ol><h3>How We’ll Learn</h3><p>Course sessions include lecture, group discussions, and hands-on FactSet terminal exercises, both in class and for homework.</p><div class="callout"><strong>Stay connected:</strong> Check email and Canvas at least once per day. Keep up with assignments, ask questions when concepts are unclear, and use the instructor’s support.</div><h3>Readings & Materials</h3><ul><li>FactSet modules in video form.</li><li>FactSet Certificate programs are free for students who access learning information through the FactSet app/web pages.</li><li>Required: FactSet Desktop Application and Microsoft Office 365.</li></ul></section>
    <section id="assessment"><h2>Evaluation & Grading</h2><p>{esc(evaluation_summary)} For the four core sections, each student or group—depending on class size—will demonstrate proficiency using the FactSet terminal to access and explain relevant financial information. Deliverables include PowerPoint submissions and verbal class presentations using the FactSet terminal.</p><div class="weights">{assess_html}</div><h3>Letter Grade Scale</h3><div class="grades">{grades_html}</div></section>
    <section id="policies"><h2>Course Policies</h2>{policy_html}</section>
    <section id="support"><h2>Student Support</h2>{support_html}<h3>Course Expectations</h3><p>For each credit hour, students are expected to spend a minimum of two hours on work outside class each week. For this one-credit course, that means a minimum of two hours each week.</p><p>Students must review the <a href="https://www.endicott.edu/academics/academic-resources-support/academic-calendar/undergraduate">Academic Calendar published by the Registrar’s Office</a> (the original syllabus also referenced <a href="http://tinyurl.com/hbmfywj">this short link</a>). Class attendance is expected through the last day of scheduled finals; students must plan accordingly.</p></section>
    <section id="schedule"><h2>Topic Outline & Timeline</h2><p>This is a rough guide. Starting in Week 1, you have something due almost every week. Changes will be announced via Canvas; check it at least once per day.</p><table><thead><tr><th>Week</th><th>Week of</th><th>Assignments Due / Weekly Class Topics</th></tr></thead><tbody>{schedule_html}</tbody></table><p><small>Fall 2026 calendar alignment: classes begin September 1; Thanksgiving recess begins after November 20; classes resume November 30; regular classes end December 11; finals run December 14–18. The Week 10 instructor-planned “No Class” date is retained from the prior syllabus pattern and should be confirmed in Canvas.</small></p></section>
    <section><h2>Subject to Change</h2><p>The syllabus sets forth my course objectives and my best estimate of what we will be able to cover during the semester. To tailor the course to the specific needs of the students, I might modify the syllabus during the semester; for this reason, it is possible that not all material will be covered and/or that additional material may be assigned.</p></section></main><footer>BUS209 · Fall 2026 · Endicott College</footer></body></html>'''
    path=OUT/'BUS209_FactSet_Financial_Concepts_Syllabus_Fall_2026.html'; path.write_text(doc,encoding='utf-8'); return path

def build_canvas_html():
    """Build a Canvas Rich Content Editor-safe body fragment.

    Canvas may sanitize page-level CSS and interactive elements, so this output
    uses semantic HTML, simple tables, and conservative inline styles only.
    """
    def esc(x): return html.escape(str(x))
    section = 'margin:0 0 24px;padding:20px;border:1px solid #d8e3e8;border-radius:10px;background:#ffffff;'
    h2 = 'margin:0 0 12px;color:#12304a;font-size:26px;line-height:1.2;'
    h3 = 'margin:20px 0 8px;color:#176b87;font-size:19px;line-height:1.25;'
    p = 'margin:0 0 12px;color:#243746;line-height:1.55;'
    table = 'width:100%;border-collapse:collapse;margin:12px 0 18px;'
    th = 'padding:9px;border:1px solid #d8e3e8;background:#12304a;color:#ffffff;text-align:left;vertical-align:top;'
    td = 'padding:9px;border:1px solid #d8e3e8;color:#243746;text-align:left;vertical-align:top;'
    overview_rows=''.join(f'<tr><th scope="row" style="{th}width:28%;">{esc(k)}</th><td style="{td}">{esc(v)}</td></tr>' for k,v in overview)
    outcome_html=''.join(f'<li style="margin:0 0 7px;">{esc(x)}</li>' for x in outcomes)
    assess_rows=''.join(f'<tr><td style="{td}">{esc(n)}</td><td style="{td}width:20%;text-align:center;"><strong>{esc(w)}</strong></td></tr>' for n,w in assessments)
    grade_rows=''.join(f'<tr><th scope="row" style="{th}width:28%;">{esc(g)}</th><td style="{td}">{esc(r)}</td></tr>' for g,r in grades)
    policy_html=''.join(f'<h3 style="{h3}">{esc(t)}</h3><p style="{p}">{esc(x)}</p>' for t,x in policies)
    support_html=''.join(f'<h3 style="{h3}">{esc(t)}</h3><p style="{p}">{esc(x)}</p>' for t,x in support)
    schedule_rows=[]
    fills={'break':'#fff8e6','presentation':'#eaf5f4','final':'#f2ecff','class':'#ffffff'}
    for wk,date,topic,kind in schedule:
        fill=fills.get(kind,'#ffffff')
        schedule_rows.append(f'<tr><td style="{td}background:{fill};text-align:center;">{wk}</td><td style="{td}background:{fill};">{esc(date)}</td><td style="{td}background:{fill};">{esc(topic)}</td></tr>')
    doc=f'''<!-- BUS209 Fall 2026 Canvas-safe syllabus fragment. Paste into the Canvas HTML editor. -->
<div style="max-width:960px;margin:0 auto;color:#243746;font-family:Arial,Helvetica,sans-serif;font-size:16px;line-height:1.55;">
  <div style="margin:0 0 24px;padding:28px 22px;background:#12304a;color:#ffffff;border-radius:12px;">
    <p style="margin:0 0 8px;color:#a8eee5;font-size:13px;font-weight:bold;letter-spacing:1px;text-transform:uppercase;">Endicott College · Curtis L. Gerrish School of Business</p>
    <h1 style="margin:0 0 8px;color:#ffffff;font-size:38px;line-height:1.1;">FactSet Financial Concepts</h1>
    <p style="margin:0;color:#ffffff;font-size:19px;">BUS209 · Fall 2026 · Course Syllabus</p>
  </div>
  <div style="{section}"><h2 style="{h2}">Course Overview</h2><table style="{table}"><tbody>{overview_rows}</tbody></table><h3 style="{h3}">Catalog Description</h3><p style="{p}">{esc(catalog)}</p></div>
  <div style="{section}"><h2 style="{h2}">Learning Outcomes</h2><p style="{p}">Upon completion of this course, students will be able to:</p><ol style="margin:0 0 16px;padding-left:28px;">{outcome_html}</ol><h3 style="{h3}">Teaching &amp; Learning Strategies</h3><p style="{p}">Course sessions include lecture, group discussions, and hands-on FactSet terminal exercises, both in class and for homework.</p><p style="{p}">In addition to the in-class sessions, it is your duty to regularly—at least once per day—check your email and Canvas to ensure you understand what you are expected to complete for this class.</p><p style="{p}">It is in your best interest to keep up with assignments. Please contact me with questions about the content. I am here to help. If you find any concepts unclear, let me know and let me help you. I want you to do your best in this course.</p><h3 style="{h3}">Readings &amp; Materials</h3><ul style="margin:0;padding-left:24px;"><li>FactSet modules in video form.</li><li>FactSet Certificate programs are free for students who access learning information through the FactSet app/web pages.</li><li>Required: FactSet Desktop Application and Microsoft Office 365.</li></ul></div>
  <div style="{section}"><h2 style="{h2}">Evaluation Methods</h2><p style="{p}">{esc(evaluation_summary)} For the four core sections, each student or group—depending on class size—will demonstrate proficiency using the FactSet terminal to access and explain relevant financial information. Deliverables include PowerPoint submissions and verbal class presentations using the FactSet terminal.</p><table style="{table}"><thead><tr><th scope="col" style="{th}">Assessment</th><th scope="col" style="{th}text-align:center;">Weight</th></tr></thead><tbody>{assess_rows}</tbody></table><h3 style="{h3}">Letter Grade Scale</h3><table style="{table}"><tbody>{grade_rows}</tbody></table></div>
  <div style="{section}"><h2 style="{h2}">Course Policies</h2>{policy_html}</div>
  <div style="{section}"><h2 style="{h2}">Student Support</h2>{support_html}<h3 style="{h3}">Course Expectations</h3><p style="{p}">For each credit hour, students are expected to spend a minimum of two hours on work outside class each week. For this one-credit course, that means a minimum of two hours each week.</p><p style="{p}">Students must review the <a href="https://www.endicott.edu/academics/academic-resources-support/academic-calendar/undergraduate" style="color:#176b87;text-decoration:underline;">Academic Calendar published by the Registrar’s Office</a> (the original syllabus also referenced <a href="http://tinyurl.com/hbmfywj" style="color:#176b87;text-decoration:underline;">this short link</a>). Class attendance is expected through the last day of scheduled finals; students must plan accordingly.</p></div>
  <div style="{section}"><h2 style="{h2}">Topic Outline &amp; Timeline</h2><p style="{p}">This is a rough guide. Starting in Week 1, you have something due almost every week. Changes will be announced via Canvas; check it at least once per day.</p><table style="{table}"><thead><tr><th scope="col" style="{th}width:10%;">Week</th><th scope="col" style="{th}width:23%;">Week of</th><th scope="col" style="{th}">Assignments Due / Weekly Class Topics</th></tr></thead><tbody>{''.join(schedule_rows)}</tbody></table><p style="margin:0;color:#5b6b78;font-size:13px;line-height:1.45;">Fall 2026 calendar alignment: classes begin September 1; Thanksgiving recess begins after November 20; classes resume November 30; regular classes end December 11; finals run December 14–18. The Week 10 instructor-planned “No Class” date should be confirmed in Canvas.</p></div>
  <div style="{section}"><h2 style="{h2}">Subject to Change Statement</h2><p style="{p}">The syllabus sets forth my course objectives and my best estimate of what we will be able to cover during the semester. To tailor the course to the specific needs of the students, I might modify the syllabus during the semester; for this reason, it is possible that not all material will be covered and/or that additional material may be assigned.</p></div>
</div>'''
    path=OUT/'BUS209_FactSet_Financial_Concepts_Syllabus_Fall_2026_CANVAS.html'
    path.write_text(doc,encoding='utf-8')
    return path

if __name__=='__main__':
    print(build_docx()); print(build_html()); print(build_canvas_html())
